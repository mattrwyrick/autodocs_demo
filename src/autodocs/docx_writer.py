import docx
from docx.shared import Inches, Pt
from django.http import HttpResponse
import urllib2
import StringIO


class New_Document(object):
    """
    Tool to enhance python-docx writing
    """

    def __init__(self):
        self.document = docx.Document()

    def get_docx_document(self):
        """
        Returns the underlying python-docx Document object

        :return: docx
        """
        return self.document

    def add_heading(self, text="", level=0, bold=False, italic=False, underline=False, font=False, size=False):
        """
        Returns an enhanced document heading

        :param text:
        :param level:
        :param bold:
        :param italic:
        :param underline:
        :return:
        :rtype: New_Header
        """
        return New_Header(self.document, text, level, bold, italic, underline, font, size)

    def add_paragraph(self, text="", bold=False, italic=False, underline=False, font=False, size=False):
        """
        Returns an enhanced document paragraph

        :param text:
        :param bold:
        :param italic:
        :param underline:
        :return:
        :rtype: New_Paragraph
        """
        return New_Paragraph(self.document, text, bold, italic, underline, font, size)

    def add_table(self, rows=0, cols=0):
        """
        Returns an enhanced document table

        :param rows:
        :param cols:
        :return:
        :rtype: New_Table
        """
        return New_Table(self.document, rows, cols)

    def add_photo_from_web(self, url, size=1):
        """
        Adds photo from the web to Document

        :param url:
        :param size:
        :return:
        """
        try:
            image_from_url = urllib2.urlopen(str(url))
            io_url = StringIO.StringIO()
            io_url.write(image_from_url.read())
            io_url.seek(0)
            self.document.add_picture(io_url, width=Inches(size))
            return True
        except:
            return False

    def create_output(self, filename="AutoDocs", pdf=False):
        """
        Creates output file for document

        :returns: output
        :rtype: StringIO / HTTPResponse
        """
        if pdf:
            content_type = 'application/pdf'
            extension = "pdf"
        else:
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            extension = "docx"

        output = StringIO.StringIO()
        self.document.save(output)
        length = output.tell()
        output.seek(0)
        response = HttpResponse(output.getvalue(), content_type=content_type)
        response['Content-Disposition'] = 'attachment; filename=%s.%s' % (filename, extension)
        response['Content-Length'] = length
        return response


    def create_file(self):
        """
        Creates output file for document
        :returns: output
        :rtype: StringIO
        """
        output = StringIO.StringIO()
        self.document.save(output)
        output.seek(0)
        return output


class New_Header(object):
    """
    Enhanced python-docx heading
    """

    def __init__(self, doc, text="", level=0, bold=False, italic=False, underline=False, font=False, size=False):
        self.header = doc.add_heading("", level)
        if text != "":
            self.write(text=str(text), bold=bold, italic=italic, underline=underline, font=font, size=size)

    def get_docx_heading(self):
        """
        Returns underlying python-docx header object

        :return:
        """
        return self.header

    def write(self, text="", bold=False, italic=False, underline=False, font=False, size=False):
        """
        Writes inline to heading

        :param text:
        :param bold:
        :param italic:
        :param underline:
        :param font:
        :param size:
        :return:
        """
        run = self.header.add_run(str(text))
        run.bold = bold
        run.italic = italic
        run.underline = underline
        if font:
            run.font.name = font
        if size:
            try:
                run.font.size = Pt(int(size))
            except:
                pass

        return text

    def write_to_new_line(self, text="", bold=False, italic=False, underline=False, font=False, size=False):
        """
        Writes to a new line in heading

        :param text:
        :param bold:
        :param italic:
        :param underline:
        :param font:
        :param size:
        :return:
        """
        return self.write(text='\n'+str(text), bold=bold, italic=italic, underline=underline, font=font, size=size)

    def add_photo_from_web(self, url, size=1):
        """
        Adds a photo from the web inline with the heading

        :param url:
        :param size: int
        :return:
        """
        try:
            image_from_url = urllib2.urlopen(str(url))
            io_url = StringIO.StringIO()
            io_url.write(image_from_url.read())
            io_url.seek(0)
            self.header.add_run().add_picture(io_url, width=Inches(size))
            return True
        except:
            return False


class New_Paragraph(object):
    """
    Enhanced python-docx paragraph
    """

    def __init__(self, doc, text="", bold=False, italic=False, underline=False, font=False, size=False):
        self.paragraph = doc.add_paragraph()
        if text != "":
            self.write(text=str(text), bold=bold, italic=italic, underline=underline, font=font, size=size)

    def get_docx_paragraph(self):
        """
        Returns underlying python-docx paragraph

        :return:
        """
        return self.paragraph

    def write(self, text="", bold=False, italic=False, underline=False, font=False, size=False):
        """
        Writes inline to paragraph

        :param text:
        :param bold:
        :param italic:
        :param underline:
        :param font:
        :param size:
        :return:
        """
        run = self.paragraph.add_run(str(text))
        run.bold = bold
        run.italic = italic
        run.underline = underline
        if font:
            run.font.name = font
        if size:
            try:
                run.font.size = Pt(int(size))
            except:
                pass

        return text

    def write_to_new_line(self, text="", bold=False, italic=False, underline=False, font=False, size=False):
        """
        Writes to a new line in paragraph

        :param text:
        :param bold:
        :param italic:
        :param underline:
        :param font:
        :param size:
        :return:
        """
        return self.write(text='\n' + str(text), bold=bold, italic=italic, underline=underline, font=font, size=size)

    def add_photo_from_web(self, url, size=1):
        """
        Adds photo from the web inline with paragraph

        :param url:
        :param size:
        :return:
        """
        try:
            image_from_url = urllib2.urlopen(str(url))
            io_url = StringIO.StringIO()
            io_url.write(image_from_url.read())
            io_url.seek(0)
            self.paragraph.add_run().add_picture(io_url, width=Inches(size))
            return True
        except:
            return False


class New_Table(object):
    """
    Enhanced python-docx table
    """

    def __init__(self, doc, rows, cols):
        self.rows = rows
        self.cols = cols
        self.table = doc.add_table(rows=rows, cols=cols)

    def get_docx_table(self):
        """
        Returns underlying python-docx table
        :return:
        """
        return self.table

    def write_to_cell(self, row, col, text):
        """
        Write to a specified cell

        :param row:
        :param col:
        :param text:
        :return:
        """
        if (row > self.rows -1) or (col > self.cols - 1):
            return False
        row = self.table.rows[row].cells
        row[col].text = str(text)
        return True

    def write_to_row_with_list(self, row, args):
        """
        Write to a specified row with a list as column arguments
        :param row:
        :param args: list of arguements
        :return:
        """
        if (row > self.rows -1) or (len(args) > self.cols):
            return False
        row = self.table.rows[row].cells
        for col in range(len(args)):
            row[col].text = str(args[col])
        return True

    def add_row_with_list(self, args):
        """
        Add a row with a list as column arguments

        :param args:
        :return:
        """
        if len(args) > self.cols:
            return False
        self.table.add_row()
        self.rows += 1
        return self.write_to_row_with_list(self.rows - 1, args)

    def write_to_row(self, row, *args):
        """
        Write to specified row with individual parameters as column arguments

        :param row:
        :param args: unspecified number column arguments
        :return:
        """
        self.write_to_row_with_list(row, args)

    def add_row(self, *args):
        """
        Add row with unspecified number of column arguements - no *args just adds row

        :param args: unspecified number of column arguments
        :return:
        """
        self.add_row_with_list(args)


class Smart_Heading(object):

    def __init__(self, text="", level=0, bold=False, italic=False, underline=False, condition=lambda x: True, font=False):
        self.text = text
        self.bold = bold
        self.italic = italic
        self.underline = underline
        self.condition = condition
        self.level = level
        self.needs_displayed = True
        self.font = font


class Smart_Headings_Displayer(object):

    def __init__(self, doc, headings=[], loop=False):
        """
        Constructor

        :param doc: New_Document
        :param headings: list of Smart_Heading objects
        :param loop: loop over headings list if end is reached
        """
        self.doc = doc
        self.loop = loop
        self.headings = headings
        self.headings.append(self.get_empty_smart_heading())
        self.current_index = 0
        self.max_index = len(self.headings) - 1 if loop else len(self.headings)

    def get_empty_smart_heading(self):
        heading = Smart_Heading(text="", level=0, bold=False, italic=False, underline=False, condition=lambda x: True)
        heading.needs_displayed = False
        return heading

    def reset_headings_display(self):
        for i in range(len(self.headings)-1):
            self.headings[i].needs_displayed = True

    def update_heading(self, data):
        """
        displays the heading the first time the condition is met. Displays the next occurring heading when the
        condition is not met.

        :param data:
        :return:
        """
        if self.loop and (self.current_index >= self.max_index):
            self.current_index = 0
            self.reset_headings_display()
            self.update_heading(data)

        heading = self.headings[self.current_index]

        if heading.needs_displayed:
            if heading.condition(data):
                self.doc.add_heading(text=heading.text, level=heading.level, bold=heading.bold,
                                     italic=heading.italic, underline=heading.underline, font=heading.font)
                heading.needs_displayed = False
            else:
                self.current_index += 1
                self.update_heading(data)

        elif not heading.condition(data):
            self.current_index += 1
            self.update_heading(data)
